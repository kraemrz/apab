import os
import re
import sys
from flask import Flask, render_template, request, jsonify, send_file, send_from_directory
from werkzeug.utils import secure_filename
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from io import BytesIO
from datetime import datetime
from database import add_inspection, get_history_for_machine
import json

SAVE_FOLDER = 'Sparade_Rapporter'
os.makedirs(SAVE_FOLDER, exist_ok=True)

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)


app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

TRANSLATIONS = {
    'sv': { 'toc_title': 'Innehåll', 'report_title': 'Inspektionsprotokoll', 'keywords': ['status', 'åtgärd', 'kommentar', 'signatur'] },
    'en': { 'toc_title': 'Table of Contents', 'report_title': 'Inspection', 'keywords': ['results', 'description', 'comment', 'signature'] }
}

def parse_filename_for_info(filename, lang):
    try:
        prefix = TRANSLATIONS[lang]['report_title']
        clean_name = filename.lower().replace('.docx', '').replace('_mall', '')
        if clean_name.startswith(prefix.lower()):
            clean_name = clean_name[len(prefix):].strip()
        
        match = re.search(r'(m\d{6})', clean_name)
        if match:
            machine = match.group(1).upper()
            customer_raw = clean_name[:match.start()]
            customer = customer_raw.replace('_', ' ').replace('-', ' ').strip().title()
            return customer if customer else "Okänd Kund", machine
    except Exception as e:
        print(f"Error parsing filename: {e}")
    return "Okänd Kund", "Okänd Maskin"

def set_cell_shade(cell, shade):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), shade)
    tcPr.append(shd)

def add_table_of_contents(document, toc_title):
    document.add_heading(toc_title, level=1)
    paragraph = document.add_paragraph(); run = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar'); fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'TOC \\o "2-2" \\h \\z \\u'
    fldChar_separate = OxmlElement('w:fldChar'); fldChar_separate.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar'); fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar_begin, instrText, fldChar_separate, fldChar_end])

def detect_language_from_doc(document):
    text_content = [cell.text.lower() for table in document.tables for row in table.rows for cell in row.cells]
    full_text = " ".join(text_content)
    return 'en' if any(keyword in full_text for keyword in TRANSLATIONS['en']['keywords']) else 'sv'

@app.route("/ping", methods=["GET", "HEAD"])
def ping_server():
    return "ok", 200 # Returnerar bara en enkel "ok" med status 200

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files: return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if not file.filename: return jsonify({"error": "No selected file"}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    try:
        doc = Document(filepath)
        detected_lang = detect_language_from_doc(doc)
        customer, machine = parse_filename_for_info(filename, detected_lang)

        history_data = get_history_for_machine(machine)
        
        blocks = []
        content_started = False

        for child in doc.element.body:
            if not content_started:
                if isinstance(child, CT_P):
                    p = Paragraph(child, doc)
                    if p.text.strip().lower().startswith('station') or re.match(r'^\d+\s+station', p.text.strip().lower()):
                        content_started = True
                
                if not content_started:
                    continue

            if isinstance(child, CT_P):
                p = Paragraph(child, doc)
                text = p.text.strip()
                if text:
                    blocks.append({"type": "paragraph", "text": text})
            elif isinstance(child, CT_Tbl):
                t = Table(child, doc)
                table_data = [[cell.text.strip() for cell in row.cells] for row in t.rows]
                blocks.append({"type": "table", "data": table_data})

        return jsonify({
            "blocks": blocks, 
            "lang": detected_lang, 
            "customer": customer, 
            "machine": machine,
            "history": history_data
            })

    
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return jsonify({"error": f"Failed to process document: {e}"}), 500
    finally:
        if os.path.exists(filepath):
            os.remove(filepath)

@app.route('/history/<machine_id>')
def get_history(machine_id):
    history = get_history_for_machine(machine_id)
    return jsonify(history)


@app.route("/export-word", methods=["POST"])
def export_to_word():
    lang = request.form.get('lang', 'sv')
    html_content = request.form.get("html", "")
    inspection_date = request.form.get('inspection_date', datetime.now().strftime('%Y-%m-%d'))
    customer = request.form.get('customer', 'Okänd Kund')
    machine = request.form.get('machine', 'Okänd Maskin')
    signature = request.form.get('signature', '') # <-- NY KOD: Hämta signaturen

    comments_json = request.form.get('comments_json', '[]')
    comments_data = json.loads(comments_json)

    soup = BeautifulSoup(html_content, "html.parser")
    document = Document()
    
    document.add_picture(resource_path("static/images/apab_logo.png"), width=Inches(3.0))
    p = document.add_paragraph(); p.add_run(TRANSLATIONS[lang]['report_title']).bold = True
    p.runs[0].font.size = Pt(24); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    
    # Ändra till 4 rader för att inkludera signaturen
    info_table = document.add_table(rows=4, cols=2) 
    info_table.columns[0].width = Inches(1.5); info_table.columns[1].width = Inches(4.5)
    
    cell_label_c = info_table.cell(0, 0).paragraphs[0]; cell_label_c.add_run('Kund:').bold = True; cell_label_c.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_table.cell(0, 1).text = customer
    cell_label_m = info_table.cell(1, 0).paragraphs[0]; cell_label_m.add_run('Maskin:').bold = True; cell_label_m.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_table.cell(1, 1).text = machine
    cell_label_d = info_table.cell(2, 0).paragraphs[0]; cell_label_d.add_run('Inspektionsdatum:').bold = True; cell_label_d.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_table.cell(2, 1).text = inspection_date
    
    # NY KOD: Lägg till signatur i infotabellen
    cell_label_s = info_table.cell(3, 0).paragraphs[0]; cell_label_s.add_run('Signatur:').bold = True; cell_label_s.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_table.cell(3, 1).text = signature # <-- Sätt signaturvärdet här

    document.add_page_break()

    add_table_of_contents(document, TRANSLATIONS[lang]['toc_title'])
    document.add_page_break()

   
    for element in soup.children:
        if not hasattr(element, 'name') or not element.name: continue
        text = element.get_text(strip=True)
        if not text: continue
        
        if element.name == 'h1': document.add_heading(text, level=2)
        elif element.name == 'h2': document.add_heading(text, level=3)
        elif element.name == 'div' and 'na-bar' in element.get('class', []):
            document.add_paragraph(f"{text} (Ej aktuell)", style='List Bullet')
        elif element.name == 'table':
            rows_data = [[cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])] for row in element.find_all('tr')]
            if rows_data and rows_data[0]:
                try:
                    doc_table = document.add_table(rows=1, cols=len(rows_data[0]), style='Table Grid')
                    header_cells = doc_table.rows[0].cells
                    for j, cell_text in enumerate(rows_data[0]):
                        header_cells[j].text = cell_text
                        set_cell_shade(header_cells[j], 'FFC000')
                    for i in range(1, len(rows_data)):
                        row_cells = doc_table.add_row().cells
                        for j, cell_text in enumerate(rows_data[i]):
                            row_cells[j].text = cell_text
                    document.add_paragraph()
                except IndexError: print("Skipping malformed table.")
    
    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    
    filename_prefix = TRANSLATIONS[lang]['report_title']
    download_name = f"{filename_prefix}_{customer.replace(' ', '_')}_{machine.replace(' ', '_')}_{inspection_date}.docx"
    try:
        add_inspection(
            customer=customer, 
            machine=machine, 
            inspection_date=inspection_date,
            comments=comments_data
            )
        print(f"INFO: Inspection added for {customer} - {machine} on {inspection_date}")
    except Exception as e:
        print(f"Error saving inspection: {e}")
    return send_file(doc_io, as_attachment=True, download_name=download_name, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route('/sw.js')
def service_worker():
    return send_from_directory('.', 'sw.js')  

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')
